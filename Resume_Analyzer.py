import pdfplumber
import docx
import spacy
from sentence_transformers import SentenceTransformer, util
from fastapi import FastAPI, UploadFile, File, Form
from io import BytesIO

# Load spaCy model for Named Entity Recognition (NER)
nlp = spacy.load("en_core_web_sm")

# Load BERT-based model for similarity matching
bert_model = SentenceTransformer('all-MiniLM-L6-v2')

app = FastAPI()

# Fix: Read PDF asynchronously
async def extract_text_from_pdf(file):
    text = ""
    file_data = await file.read()  # Fix: Read file asynchronously
    with pdfplumber.open(BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip() if text.strip() else "No readable text found in PDF."

# Fix: Read DOCX asynchronously
async def extract_text_from_docx(file):
    file_data = await file.read()  # Fix: Read file asynchronously
    doc = docx.Document(BytesIO(file_data))
    return "\n".join([para.text for para in doc.paragraphs]) if doc.paragraphs else "No text found in DOCX file."

# NLP Processing
def extract_keywords(text):
    doc = nlp(text)
    return [token.lemma_ for token in doc if token.pos_ in ["NOUN", "PROPN"]]

def calculate_similarity(resume_text, job_desc):
    if not resume_text.strip():  # Handle empty resume text
        return 0.0
    resume_embedding = bert_model.encode(resume_text, convert_to_tensor=True)
    job_embedding = bert_model.encode(job_desc, convert_to_tensor=True)
    similarity_score = util.pytorch_cos_sim(resume_embedding, job_embedding).item()
    return round(similarity_score * 100, 2)

# Fix: Ensure async functions are called properly
@app.post("/analyze_resume/")
async def analyze_resume(file: UploadFile = File(...), job_description: str = Form(...)):
    try:
        print(f"Received file: {file.filename}")
        print(f"Job description: {job_description}")

        # Extract text properly
        if file.filename.endswith(".pdf"):
            resume_text = await extract_text_from_pdf(file)
        elif file.filename.endswith(".docx"):
            resume_text = await extract_text_from_docx(file)
        else:
            return {"error": "Unsupported file format. Upload a PDF or DOCX file."}

        print(f"Extracted resume text (first 500 chars): {resume_text[:500]}")

        # NLP Analysis
        keywords = extract_keywords(resume_text)
        print(f"Extracted Keywords: {keywords}")

        similarity = calculate_similarity(resume_text, job_description)
        print(f"Match Score: {similarity}%")

        return {
            "resume_keywords": keywords,
            "match_score": similarity,
            "suggestion": "Add missing skills from job description to improve your match."
        }

    except Exception as e:
        print(f"Internal Server Error: {e}")  # Debug print
        return {"error": f"Internal server error: {str(e)}"}
